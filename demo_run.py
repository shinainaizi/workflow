#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Demo Mode - No API Key required
演示模式：无需API Key，使用规则引擎从方案和回忆记录中提取信息并生成台账和报告
"""

import os
import sys
import re
from pathlib import Path
from datetime import datetime
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def safe_print(msg):
    """Safe print that handles encoding issues"""
    try:
        print(msg)
    except UnicodeEncodeError:
        safe_msg = msg.encode('ascii', 'ignore').decode('ascii')
        if safe_msg.strip():
            print(safe_msg)
        else:
            print("[Non-ASCII content]")


def formalize_text(text, field_type='general'):
    """
    将口语化文本转换为书面化表达（规则引擎基础版）
    field_type: 'highlights' | 'problems' | 'improvements' | 'general'
    注意：此为演示模式的规则引擎，做基础替换；接入AI API后由大模型智能改写
    """
    if not text or not text.strip():
        return text

    lines = [l.strip() for l in text.split('\n') if l.strip()]
    formalized = []

    for line in lines:
        result = line

        # 去掉引号，保留内容
        result = re.sub(r'["\u201c\u201d]([^\u201c\u201d"]*)["\u201c\u201d"]', r'\1', result)

        # 去掉括号中的口语补充（如"叫豆豆"）
        result = re.sub(r'（叫[^）]*）', '', result)
        result = re.sub(r'\(叫[^)]*\)', '', result)

        # 去掉"有个"前缀
        result = re.sub(r'^有个', '部分', result)
        result = re.sub(r'，有个', '，部分', result)

        # 去掉残留的口语连接（"说"在句首/逗号后单独出现时替换）
        result = re.sub(r'^说，?', '', result)
        result = re.sub(r'，说，', '，', result)

        # 去掉时间残留（如"11:25才结束"）
        result = re.sub(r'\d{1,2}:\d{2}才结束[，。]?', '', result)
        result = re.sub(r'\d{1,2}:\d{2}结束[，。]?', '', result)

        # 口语→书面映射（按长度降序排列，优先匹配长词）
        colloquial_map = [
            # 亮点类
            (r'宝宝最喜欢(.+?)，排队抢着玩', r'\1环节受到幼儿欢迎，参与积极性高'),
            (r'排队抢着玩', '积极参与'),
            (r'宝宝最喜欢', '幼儿对相应环节表现出浓厚兴趣'),
            (r'有个宝宝[^，。]*特别活跃，帮其他小朋友找东西', '部分幼儿表现出较强的互助意识，主动协助同伴完成活动任务'),
            (r'有个宝宝[^，。]*，', '部分幼儿'),
            (r'特别活跃', '表现积极'),
            (r'帮其他小朋友', '主动协助同伴'),
            (r'家长反馈说', '家长反馈表明'),
            (r'希望以后多办', '期望后续持续开展'),
            (r'这样的活动很有意义', '该活动具有较强的教育意义'),
            (r'很有意义', '具有教育意义'),
            (r'从来没见过这么有意思的活动', '活动形式新颖，获得高度认可'),
            (r'有意思', '具有趣味性'),
            (r'老奶奶带着孙子来的', '隔代家庭成员参与'),
            (r'隔代家庭成员参与，说活动形式新颖，获得高度认可', '隔代家庭成员积极参与，对活动形式给予高度评价'),
            (r'来了', '到场参与'),
            (r'宝宝很投入', '幼儿专注度较高'),
            (r'很可爱', '形象生动'),
            # 问题类
            (r'准备的(.+?)不够', r'\1储备不足'),
            (r'后面来的家庭没有拿到', '后到场家庭未能领取'),
            (r'不够', '数量不足'),
            (r'只有(\d+)个', r'仅有\1个'),
            (r'需要轮流使用', '需轮换使用'),
            (r'等待时间有点长', '等待时间较长'),
            (r'有点小', '音量不足'),
            (r'听不清', '难以清晰接收'),
            (r'比预计长了(\d+)分钟', r'超出预计\1分钟'),
            (r'比预计多(\d+)分钟', r'超出预计\1分钟'),
            (r'原计划.*?实际.*?才结束', '计划与实际执行存在偏差'),
            (r'原计划.*?实际', '计划与实际执行存在偏差'),
            # 改进类
            (r'下次多准备些', '后续应增加'),
            (r'多准备', '增加配备'),
            (r'提前测试', '提前调试'),
            (r'严格控制时间', '加强时间管控'),
            (r'每个环节设闹钟提醒', '各环节设置时间提醒'),
            (r'可以考虑分批次进行', '建议实行分批次开展'),
            (r'避免拥挤', '降低现场密度'),
            # 通用
            (r'很好', '良好'),
            (r'很多', '较多'),
            (r'大家都', '参与者普遍'),
            (r'大家', '参与者'),
            (r'最有意思的环节', '最具互动性的环节'),
            (r'很有', '具有较强的'),
        ]

        for pattern, replacement in colloquial_map:
            result = re.sub(pattern, replacement, result)

        if result.strip():
            formalized.append(result.strip())

    return '；'.join(formalized) if formalized else text


def extract_from_scheme(scheme_path):
    """从策划方案Word文档中提取结构化信息"""
    doc = docx.Document(scheme_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # 将段落组装成文本，保留换行
    full_text = '\n'.join(paragraphs)

    data = {
        'activity_theme': '',
        'activity_date': '',
        'activity_time': '',
        'activity_location': '',
        'participants': '',
        'expected_participants': '',
        'actual_participants': '',
        'activity_goal': '',
        'activity_content': '',
        'activity_flow': [],
        'highlights': '',
        'problems': '',
        'improvements': '',
        'staff': '社区工作人员',
        'notes': ''
    }

    # ===== 提取各字段 =====
    # 文档格式为"一、活动目的\n内容"，即标题行后面跟内容行
    # 也兼容"活动主题：内容"格式

    def extract_field(field_names, text, paragraphs):
        """通用字段提取：支持'标题：内容'和'X、标题\n内容'两种格式"""
        for name in field_names:
            # 格式1: "活动主题：内容" 或 "活动主题: 内容"
            pattern1 = re.compile(r'(?:' + name + r')\s*[:：]\s*(.+)')
            match1 = pattern1.search(text)
            if match1:
                return match1.group(1).strip()

            # 格式2: "X、活动主题\n内容"（标题行，下一行是内容）
            for i, p in enumerate(paragraphs):
                if re.search(r'[\d一二三四五六七八九十]+[、.\s]' + name, p):
                    if i + 1 < len(paragraphs):
                        return paragraphs[i + 1].strip()
        return ''

    # 活动主题
    data['activity_theme'] = extract_field(
        ['活动主题', '主题'], full_text, paragraphs
    )

    # 活动时间（含日期和时间段）
    time_field = extract_field(['活动时间', '时间'], full_text, paragraphs)
    if time_field:
        # 提取日期：如 "6 月 25 日" 或 "6月25日"
        date_match = re.search(r'(\d{1,2})\s*月\s*(\d{1,2})\s*日', time_field)
        if date_match:
            year = datetime.now().year
            data['activity_date'] = f"{year}年{int(date_match.group(1))}月{int(date_match.group(2))}日"

        # 也检查回忆记录中的完整日期
        # 提取时间段：如 "10:45-11:15"
        time_match = re.search(r'(\d{1,2}:\d{2})\s*[-~～]\s*(\d{1,2}:\d{2})', time_field)
        if time_match:
            data['activity_time'] = f"{time_match.group(1)}-{time_match.group(2)}"

    # 活动地点
    data['activity_location'] = extract_field(
        ['活动地点', '地点'], full_text, paragraphs
    )

    # 参与对象
    data['participants'] = extract_field(
        ['参与对象', '对象'], full_text, paragraphs
    )

    # 从参与对象中提取预期人数
    if data['participants']:
        expected_match = re.search(r'(\d+)\s*组', data['participants'])
        if expected_match:
            data['expected_participants'] = f"{expected_match.group(1)}组家庭"

    # 活动目的
    data['activity_goal'] = extract_field(
        ['活动目的', '目的'], full_text, paragraphs
    )

    # 活动流程（提取各环节标题）
    flow_items = []
    for p in paragraphs:
        # 匹配 "环节一：xxx" 或 "环节一：xxx（3 分钟）"
        flow_match = re.match(r'环节[一二三四五六七八九十]+[：:]\s*(.+)', p)
        if flow_match:
            flow_items.append(flow_match.group(1).strip())
        # 也匹配 "X. xxx" 或 "X、xxx" 格式的环节
        elif re.match(r'[\d]+[.、]\s*(.+)', p) and ('分钟' in p or '环节' in p):
            flow_match2 = re.match(r'[\d]+[.、]\s*(.+)', p)
            if flow_match2:
                flow_items.append(flow_match2.group(1).strip())

    # 如果没找到环节，尝试匹配带时间标识的段落
    if not flow_items:
        for p in paragraphs:
            if re.search(r'\(\d+\s*分钟\)|（\d+\s*分钟）', p):
                # 去掉前面的序号
                clean = re.sub(r'^[\d一二三四五六七八九十]+[、.\s]*', '', p)
                if clean:
                    flow_items.append(clean.strip())

    data['activity_flow'] = flow_items[:6]  # 最多保留6个环节

    # 注意事项
    notes_idx = -1
    for i, p in enumerate(paragraphs):
        if re.search(r'[\d一二三四五六七八九十]+[、.\s]*注意事项', p):
            notes_idx = i
            break
    if notes_idx >= 0:
        note_lines = []
        for p in paragraphs[notes_idx + 1:]:
            if re.match(r'[\d一二三四五六七八九十]+[、.]', p):
                break  # 遇到下一个标题，停止
            if p.strip():
                note_lines.append(p.strip())
        data['notes'] = '\n'.join(note_lines[:4])

    return data


def extract_from_notes(notes_path, base_data):
    """从回忆记录txt中提取补充信息"""
    with open(notes_path, 'r', encoding='utf-8') as f:
        notes_text = f.read()

    # 提取实际参与人数
    actual_match = re.search(r'(?:实际到场|实际|到场)\s*(\d+)\s*组', notes_text)
    if not actual_match:
        actual_match = re.search(r'(\d+)\s*组\s*家庭', notes_text)
    if actual_match:
        base_data['actual_participants'] = f"{actual_match.group(1)}组家庭"

    # 从回忆记录中提取完整日期（方案里可能只有月日）
    date_match = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', notes_text)
    if date_match:
        base_data['activity_date'] = f"{date_match.group(1)}年{int(date_match.group(2))}月{int(date_match.group(3))}日"

    # 提取活动亮点
    highlights_match = re.search(
        r'(?:活动亮点|亮点)\s*[:：]?\s*([\s\S]+?)(?=\n\n|\n遇到的问题|\n问题|\n改进|\n现场照片|\n照片)',
        notes_text
    )
    if highlights_match:
        highlights = highlights_match.group(1).strip()
        # 清理每行前面的 - 或数字前缀
        highlights = re.sub(r'^\s*[-\d][.、)]\s*', '', highlights, flags=re.MULTILINE)
        highlights = re.sub(r'^\s*-\s+', '', highlights, flags=re.MULTILINE)
        highlights = re.sub(r'^\s*\d+[.、)]\s*', '', highlights, flags=re.MULTILINE)
        base_data['highlights'] = highlights

    # 提取问题
    problems_match = re.search(
        r'(?:遇到的问题|问题|Issues|Problems)\s*[:：]?\s*([\s\S]+?)(?=\n\n|\n改进|\n建议|\n现场照片|\n照片)',
        notes_text
    )
    if problems_match:
        problems = problems_match.group(1).strip()
        problems = re.sub(r'^\s*[-\d][.、)]\s*', '', problems, flags=re.MULTILINE)
        problems = re.sub(r'^\s*-\s+', '', problems, flags=re.MULTILINE)
        problems = re.sub(r'^\s*\d+[.、)]\s*', '', problems, flags=re.MULTILINE)
        base_data['problems'] = problems

    # 提取改进建议
    improvements_match = re.search(
        r'(?:改进建议|建议|Improvements|Suggestions)\s*[:：]?\s*([\s\S]+?)(?=\n\n|\n现场照片|\n照片|\Z)',
        notes_text
    )
    if improvements_match:
        improvements = improvements_match.group(1).strip()
        improvements = re.sub(r'^\s*[-\d][.、)]\s*', '', improvements, flags=re.MULTILINE)
        improvements = re.sub(r'^\s*-\s+', '', improvements, flags=re.MULTILINE)
        improvements = re.sub(r'^\s*\d+[.、)]\s*', '', improvements, flags=re.MULTILINE)
        base_data['improvements'] = improvements

    # 生成活动内容描述
    base_data['activity_content'] = generate_content_description(base_data)

    return base_data


def generate_content_description(data):
    """生成活动内容描述（中文）"""
    lines = []

    date_str = data.get('activity_date', '')
    time_str = data.get('activity_time', '')
    location = data.get('activity_location', '社区')
    theme = data.get('activity_theme', '')
    actual = data.get('actual_participants', '')
    goal = data.get('activity_goal', '')
    flow = data.get('activity_flow', [])

    # 开头段
    time_part = f" {time_str}" if time_str else ""
    lines.append(f"{date_str}{time_part}，{location}开展「{theme}」活动。")

    if actual:
        lines.append(f"本次活动共吸引{actual}参与。")
    elif data.get('expected_participants'):
        lines.append(f"本次活动共吸引{data['expected_participants']}参与。")

    lines.append("")

    # 活动目的
    if goal:
        lines.append(f"活动旨在{goal}")

    lines.append("")

    # 活动流程
    if flow:
        lines.append("活动依次开展以下环节：")
        for i, item in enumerate(flow, 1):
            lines.append(f"  {i}. {item}")

    lines.append("")

    lines.append("现场氛围热烈欢快，获得参与家庭一致好评。社区将继续推出多样化活动，丰富居民文化生活。")

    return '\n'.join(lines)


def generate_ledger(data, output_dir='outputs/ledger'):
    """生成活动台账Word文档"""
    os.makedirs(output_dir, exist_ok=True)

    doc = docx.Document()

    # 页边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('社区活动台账')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

    doc.add_paragraph()  # 空行

    # 表格
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    # 列宽
    table.columns[0].width = Cm(3.5)
    table.columns[1].width = Cm(12)

    # 第1行：时间
    row0 = table.rows[0].cells
    row0[0].text = '时间'
    row0[1].text = f"{data.get('activity_date', '')} {data.get('activity_time', '')}"

    # 第2行：活动主题
    row1 = table.rows[1].cells
    row1[0].text = '活动主题'
    row1[1].text = data.get('activity_theme', '')

    # 第3行：参与人员
    row2 = table.rows[2].cells
    row2[0].text = '参与人员'
    participants_text = data.get('participants', '')
    actual = data.get('actual_participants', '')
    expected = data.get('expected_participants', '')
    if actual or expected:
        participants_text += f"（预期{expected}，实际{actual}）"
    row2[1].text = participants_text

    # 第4行：活动地点
    row3 = table.rows[3].cells
    row3[0].text = '活动地点'
    row3[1].text = data.get('activity_location', '')

    # 第5行：活动内容（合并单元格）
    row4 = table.rows[4].cells
    row4[0].merge(row4[1])
    row4[0].text = '活动内容：'

    content = data.get('activity_content', '')
    if content:
        content_para = row4[0].paragraphs[0]
        content_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for i, line in enumerate(content.split('\n')):
            if i == 0:
                run = content_para.add_run(line)
            else:
                content_para.add_run('\n')
                run = content_para.add_run(line)
            run.font.size = Pt(11)
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # 第6行：备注（合并单元格）
    row5 = table.rows[5].cells
    row5[0].merge(row5[1])
    row5[0].text = '备注：'

    notes_parts = []

    if data.get('highlights'):
        highlights = formalize_text(data['highlights'], 'highlights')
        notes_parts.append(f"亮点：{highlights[:200]}")

    if data.get('problems'):
        problems = formalize_text(data['problems'], 'problems')
        notes_parts.append(f"问题：{problems[:200]}")

    if data.get('improvements'):
        improvements = formalize_text(data['improvements'], 'improvements')
        notes_parts.append(f"改进：{improvements[:200]}")

    if notes_parts:
        notes_para = row5[0].paragraphs[0]
        notes_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for i, line in enumerate(notes_parts):
            if i == 0:
                run = notes_para.add_run(line)
            else:
                notes_para.add_run('\n')
                run = notes_para.add_run(line)
            run.font.size = Pt(10)
            run.font.name = 'SimSun'
            run.font.color.rgb = RGBColor(100, 100, 100)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # 格式化表格
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    if not run.font.size:
                        run.font.size = Pt(12)
                    if not run.font.name:
                        run.font.name = 'SimSun'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # 保存
    theme = data.get('activity_theme', '台账')
    # 清理文件名中的非法字符
    safe_theme = re.sub(r'[/\\:*?"<>|]', '-', theme)
    date_safe = data.get('activity_date', datetime.now().strftime('%Y%m%d'))
    date_safe = re.sub(r'[年月日\s]', '', date_safe)
    # 补零：如 "2026625" → "20260625"
    # 先尝试匹配 "年年年年月月日日" 格式（8位）
    date_match = re.match(r'(\d{4})(\d{2})(\d{2})$', date_safe)
    if not date_match:
        # 再尝试 "年年年年月日" 格式（7位），月份1位
        date_match = re.match(r'(\d{4})(\d{1})(\d{1,2})$', date_safe)
    if not date_match:
        # 尝试 "年年年年月月日" 格式（7位），日期1位
        date_match = re.match(r'(\d{4})(\d{2})(\d{1})$', date_safe)
    if date_match:
        date_safe = f"{date_match.group(1)}{int(date_match.group(2)):02d}{int(date_match.group(3)):02d}"
    filename = f"{safe_theme}_{date_safe}.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)

    return output_path


def generate_report(activities, output_dir='outputs/report'):
    """生成汇总报告Word文档"""
    os.makedirs(output_dir, exist_ok=True)

    doc = docx.Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('社区活动汇总报告')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'（{datetime.now().strftime("%Y年%m月")}）')
    run.font.size = Pt(14)
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    doc.add_paragraph()  # 空行

    # 一、总体概况
    doc.add_heading('一、总体概况', level=1)
    total = len(activities)

    # 计算总参与人数
    total_participants = 0
    for a in activities:
        actual = str(a.get('actual_participants', '0'))
        num_match = re.search(r'\d+', actual)
        if num_match:
            total_participants += int(num_match.group())

    overview = f"本周期内共举办社区活动{total}场，累计参与人数约{total_participants}人次。"
    doc.add_paragraph(overview)
    doc.add_paragraph(
        "活动涵盖感统训练、亲子互动、角色扮演等多种类型，"
        "服务婴幼儿、家庭及社区居民群体。各项活动按计划顺利开展，现场氛围良好，参与者反馈积极。"
    )

    # 二、活动详情
    doc.add_heading('二、活动详情', level=1)
    for i, activity in enumerate(activities, 1):
        doc.add_heading(f'{i}. {activity.get("activity_theme", "未命名活动")}', level=2)

        info_lines = [
            f'时间：{activity.get("activity_date", "")} {activity.get("activity_time", "")}',
            f'地点：{activity.get("activity_location", "")}',
            f'参与人员：{activity.get("participants", "")}',
            f'实际参与：{activity.get("actual_participants", "未知")}',
        ]
        for line in info_lines:
            p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Cm(0.5)

        content = activity.get('activity_content', '')
        if content:
            p = doc.add_paragraph('活动内容：')
            p.paragraph_format.left_indent = Cm(0.5)
            # 截取前300字
            display_content = content[:300] + '...' if len(content) > 300 else content
            p = doc.add_paragraph(display_content)
            p.paragraph_format.left_indent = Cm(0.5)

        if activity.get('highlights'):
            highlights_text = formalize_text(activity['highlights'], 'highlights')
            p = doc.add_paragraph(f'活动亮点：\n{highlights_text[:300]}')
            p.paragraph_format.left_indent = Cm(0.5)

    # 三、复盘与优化
    doc.add_heading('三、复盘与优化', level=1)

    all_problems = []
    all_improvements = []
    for activity in activities:
        if activity.get('problems'):
            all_problems.append(formalize_text(activity['problems'], 'problems'))
        if activity.get('improvements'):
            all_improvements.append(formalize_text(activity['improvements'], 'improvements'))

    if all_problems:
        doc.add_heading('（一）存在问题', level=2)
        for i, problem in enumerate(all_problems[:5], 1):
            p = doc.add_paragraph(f'{i}. {problem[:300]}')
            p.paragraph_format.left_indent = Cm(0.5)

    if all_improvements:
        doc.add_heading('（二）改进方向', level=2)
        for i, improvement in enumerate(all_improvements[:5], 1):
            p = doc.add_paragraph(f'{i}. {improvement[:300]}')
            p.paragraph_format.left_indent = Cm(0.5)

    if not all_problems and not all_improvements:
        doc.add_paragraph("各项活动开展顺利，建议持续关注参与者反馈，不断优化活动流程。")

    # 四、下一步计划
    doc.add_heading('四、下一步计划', level=1)
    plans = [
        "1. 持续优化活动流程，提升参与者体验",
        "2. 根据居民反馈丰富活动内容",
        "3. 加强活动宣传，扩大参与覆盖面",
        "4. 完善活动档案管理，建立长效机制"
    ]
    for plan in plans:
        p = doc.add_paragraph(plan)
        p.paragraph_format.left_indent = Cm(0.5)

    # 保存
    filename = f'汇总报告_{datetime.now().strftime("%Y%m%d")}.docx'
    report_path = os.path.join(output_dir, filename)
    doc.save(report_path)

    return report_path


def main():
    """主流程"""
    log_path = 'run_log.txt'
    log_lines = []

    log_lines.append("=" * 60)
    log_lines.append("社区活动台账自动化工作流 - 演示模式")
    log_lines.append("=" * 60)

    # 文件路径
    scheme_path = 'examples/scheme.docx'
    notes_path = 'examples/回忆记录.txt'

    # 检查文件是否存在
    if not os.path.exists(scheme_path):
        log_lines.append(f"错误：找不到方案文件 {scheme_path}")
        with open(log_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(log_lines))
        sys.exit(1)

    # 步骤1：从方案提取信息
    log_lines.append("")
    log_lines.append("[1/4] 正在从策划方案中提取信息...")
    data = extract_from_scheme(scheme_path)
    log_lines.append(f"  主题：{data['activity_theme']}")
    log_lines.append(f"  日期：{data['activity_date']}")
    log_lines.append(f"  时间：{data['activity_time']}")
    log_lines.append(f"  地点：{data['activity_location']}")
    log_lines.append(f"  对象：{data['participants']}")

    # 步骤2：从回忆记录补充信息
    if os.path.exists(notes_path):
        log_lines.append("")
        log_lines.append("[2/4] 正在从回忆记录中补充信息...")
        data = extract_from_notes(notes_path, data)
        log_lines.append(f"  实际参与：{data['actual_participants']}")
        log_lines.append(f"  亮点：{data['highlights'][:50]}..." if data['highlights'] else "  亮点：无")
    else:
        log_lines.append("")
        log_lines.append("[2/4] 未找到回忆记录，跳过")
        data['activity_content'] = generate_content_description(data)

    # 步骤3：生成台账
    log_lines.append("")
    log_lines.append("[3/4] 正在生成活动台账...")
    ledger_path = generate_ledger(data)
    log_lines.append(f"  已生成：{ledger_path}")

    # 步骤4：生成汇总报告
    log_lines.append("")
    log_lines.append("[4/4] 正在生成汇总报告...")
    report_path = generate_report([data])
    log_lines.append(f"  已生成：{report_path}")

    log_lines.append("")
    log_lines.append("=" * 60)
    log_lines.append("完成！")
    log_lines.append(f"台账文件：{ledger_path}")
    log_lines.append(f"报告文件：{report_path}")
    log_lines.append("=" * 60)

    # 写入日志文件
    with open(log_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(log_lines))

    # 同时打印到控制台
    for line in log_lines:
        safe_print(line)


if __name__ == '__main__':
    main()
