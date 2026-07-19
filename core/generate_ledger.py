#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
台账生成模块 - 根据提取的活动信息生成Word格式台账
"""

import os
from pathlib import Path
from typing import Dict, Any
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

class LedgerGenerator:
    """活动台账生成器"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.output_dir = config.get('output', {}).get('ledger_dir', 'outputs/ledger')
        os.makedirs(self.output_dir, exist_ok=True)
        
        # 加载模板（如果有）
        template_path = config.get('output', {}).get('ledger_template', '')
        self.template = None
        if template_path and os.path.exists(template_path):
            self.template = Document(template_path)
    
    def generate(self, activity_data: Dict[str, Any]) -> str:
        """
        生成活动台账
        
        Args:
            activity_data: 活动结构化数据
            
        Returns:
            生成的台账文件路径
        """
        # 创建新文档或使用模板
        if self.template:
            doc = docx.Document(self.template)
        else:
            doc = self._create_default_template()
        
        # 填充数据
        self._fill_ledger(doc, activity_data)
        
        # 保存文件
        filename = self._generate_filename(activity_data)
        output_path = os.path.join(self.output_dir, filename)
        doc.save(output_path)
        
        return output_path
    
    def _create_default_template(self):
        """创建默认台账模板"""
        doc = docx.Document()
        
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        return doc
    
    def _fill_ledger(self, doc: Document, data: Dict[str, Any]):
        """填充台账内容"""
        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('申沁社区活动台账')
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 空行
        doc.add_paragraph()
        
        # 创建表格（5行2列）
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        # 设置列宽
        table.columns[0].width = Cm(3)
        table.columns[1].width = Cm(12)
        
        # 填充表格数据
        # 行1: 时间 | 地点
        row1 = table.rows[0].cells
        row1[0].text = '时间'
        row1[1].text = self._format_date_time(data)
        
        # 行2: 活动主题
        row2 = table.rows[1].cells
        row2[0].text = '活动主题'
        row2[1].text = data.get('activity_theme', '')
        
        # 行3: 参加人员
        row3 = table.rows[2].cells
        row3[0].text = '参加人员'
        row3[1].text = data.get('participants', '社区居民')
        
        # 行4: 活动内容（合并单元格）
        row4 = table.rows[3].cells
        row4[0].text = '活动内容'
        row4[0].merge(row4[1])  # 合并两列
        
        # 设置活动内容文本
        content = data.get('activity_content', '')
        row4[0].text = ''  # 清空合并后的文本
        
        # 添加活动内容段落
        content_para = row4[0].paragraphs[0]
        content_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 处理换行
        lines = content.split('\n')
        for i, line in enumerate(lines):
            if i == 0:
                run = content_para.add_run(line)
            else:
                content_para.add_run('\n')
                run = content_para.add_run(line)
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 行5: 备注（合并单元格）
        row5 = table.rows[4].cells
        row5[0].text = '备注'
        row5[0].merge(row5[1])
        
        # 添加备注信息
        notes = self._generate_notes(data)
        row5[0].text = ''
        notes_para = row5[0].paragraphs[0]
        notes_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for i, line in enumerate(notes.split('\n')):
            if i == 0:
                run = notes_para.add_run(line)
            else:
                notes_para.add_run('\n')
                run = notes_para.add_run(line)
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run.font.color.rgb = RGBColor(128, 128, 128)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 设置表格样式
        self._format_table(table)
    
    def _format_table(self, table):
        """格式化表格样式"""
        for row in table.rows:
            for cell in row.cells:
                # 设置单元格内字体
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
                # 设置单元格垂直居中
                cell.vertical_alignment = 1  # CENTER
    
    def _format_date_time(self, data: Dict[str, Any]) -> str:
        """格式化日期时间"""
        date = data.get('activity_date', '')
        time = data.get('activity_time', '')
        
        # 尝试解析日期
        if date:
            try:
                # 处理 "6月25日" 格式
                if '月' in date and '日' in date:
                    import re
                    match = re.search(r'(\d{1,2})月(\d{1,2})日', date)
                    if match:
                        month = int(match.group(1))
                        day = int(match.group(2))
                        year = datetime.now().year
                        date = f"{year}.{month}.{day}"
                
                # 处理 "2026.6.26" 格式
                if '.' in date:
                    parts = date.split('.')
                    if len(parts) == 3:
                        date = f"{parts[0]}.{parts[1]}.{parts[2]}"
                    elif len(parts) == 2:
                        year = datetime.now().year
                        date = f"{year}.{parts[0]}.{parts[1]}"
            except:
                pass
        
        # 组合日期和时间
        if date and time:
            return f"{date} {time}"
        elif date:
            return date
        elif time:
            return time
        else:
            return datetime.now().strftime('%Y.%m.%d')
    
    def _generate_notes(self, data: Dict[str, Any]) -> str:
        """生成备注信息"""
        notes_parts = []
        
        # 预期 vs 实际参与人数
        expected = data.get('expected_participants', '')
        actual = data.get('actual_participants', data.get('estimated_participants', ''))
        if expected or actual:
            notes_parts.append(f"预期参与人数: {expected or '未填写'}  实际参与人数: {actual or '未统计'}")
        
        # 工作人员
        staff = data.get('staff', '')
        if staff:
            notes_parts.append(f"工作人员: {staff}")
        
        # 活动亮点
        highlights = data.get('highlights', '')
        if highlights:
            notes_parts.append(f"活动亮点: {highlights}")
        
        # 存在问题
        problems = data.get('problems', '')
        if problems:
            notes_parts.append(f"存在问题: {problems}")
        
        # 改进建议
        improvements = data.get('improvements', '')
        if improvements:
            notes_parts.append(f"改进建议: {improvements}")
        
        # 照片数量
        photos_count = data.get('photos_count', 0)
        if photos_count > 0:
            notes_parts.append(f"现场照片: {photos_count}张")
        
        # 注意事项
        activity_notes = data.get('notes', '')
        if activity_notes:
            notes_parts.append(f"注意事项: {activity_notes}")
        
        return '\n'.join(notes_parts) if notes_parts else '无备注'
    
    def _generate_filename(self, data: Dict[str, Any]) -> str:
        """生成文件名"""
        theme = data.get('activity_theme', '活动台账')
        # 清理非法字符
        theme = theme.replace('/', '-').replace('\\', '-').replace(':', '-')
        
        date = data.get('activity_date', '')
        if not date:
            date = datetime.now().strftime('%Y%m%d')
        
        return f"{theme}_{date}.docx"
