#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
汇总报告生成模块 - 生成定期汇总报告
"""

import os
import json
from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime, timedelta
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from llm_client import LLMClient

class ReportGenerator:
    """汇总报告生成器"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.output_dir = config.get('output', {}).get('report_dir', 'outputs/report')
        os.makedirs(self.output_dir, exist_ok=True)
        
        # 数据存储文件
        self.data_file = os.path.join(self.output_dir, '.activities_data.json')
        
        # LLM客户端
        self.llm = LLMClient(config)
    
    def update_report(self, activity_data: Dict[str, Any], ledger_path: str) -> str:
        """
        更新汇总报告
        
        Args:
            activity_data: 活动数据
            ledger_path: 台账文件路径
            
        Returns:
            报告文件路径
        """
        # 加载已有数据
        activities = self._load_activities()
        
        # 添加新活动
        activity_data['ledger_path'] = ledger_path
        activity_data['timestamp'] = datetime.now().isoformat()
        activities.append(activity_data)
        
        # 保存数据
        self._save_activities(activities)
        
        # 生成报告
        report_path = self._generate_report(activities)
        
        return report_path
    
    def _load_activities(self) -> List[Dict[str, Any]]:
        """加载活动数据"""
        if os.path.exists(self.data_file):
            with open(self.data_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return []
    
    def _save_activities(self, activities: List[Dict[str, Any]]):
        """保存活动数据"""
        with open(self.data_file, 'w', encoding='utf-8') as f:
            json.dump(activities, f, ensure_ascii=False, indent=2)
    
    def _generate_report(self, activities: List[Dict[str, Any]]) -> str:
        """生成汇总报告"""
        doc = Document()
        
        # 设置页面
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('社区活动工作汇总报告')
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 副标题
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        period = self._get_report_period(activities)
        run = subtitle.add_run(f'（{period}）')
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        doc.add_paragraph()
        
        # 一、总体概况
        doc.add_heading('一、总体概况', level=1)
        overview = self._generate_overview(activities)
        doc.add_paragraph(overview)
        
        # 二、活动统计
        doc.add_heading('二、活动统计', level=1)
        stats = self._generate_statistics(activities)
        doc.add_paragraph(stats)
        
        # 三、活动详情
        doc.add_heading('三、活动详情', level=1)
        for i, activity in enumerate(activities, 1):
            doc.add_heading(f'{i}. {activity.get("activity_theme", "未命名活动")}', level=2)
            
            # 基本信息
            info_lines = []
            info_lines.append(f'时间：{self._format_date(activity.get("activity_date", ""))} {activity.get("activity_time", "")}')
            info_lines.append(f'地点：{activity.get("activity_location", "")}')
            info_lines.append(f'参与人数：{activity.get("actual_participants", activity.get("estimated_participants", "未知"))}')
            info_lines.append(f'参与对象：{activity.get("participants", "")}')
            
            for line in info_lines:
                p = doc.add_paragraph(line)
                p.paragraph_format.left_indent = Cm(0.5)
            
            # 活动内容摘要
            content = activity.get('activity_content', '')
            if content:
                p = doc.add_paragraph('内容摘要：')
                p.paragraph_format.left_indent = Cm(0.5)
                p = doc.add_paragraph(content[:200] + '...' if len(content) > 200 else content)
                p.paragraph_format.left_indent = Cm(0.5)
            
            # 亮点
            highlights = activity.get('highlights', '')
            if highlights:
                p = doc.add_paragraph(f'亮点：{highlights}')
                p.paragraph_format.left_indent = Cm(0.5)
            
            # 问题
            problems = activity.get('problems', '')
            if problems:
                p = doc.add_paragraph(f'问题：{problems}')
                p.paragraph_format.left_indent = Cm(0.5)
        
        # 四、复盘与优化建议
        doc.add_heading('四、复盘与优化建议', level=1)
        analysis = self._generate_analysis(activities)
        doc.add_paragraph(analysis)
        
        # 五、下一步计划
        doc.add_heading('五、下一步计划', level=1)
        plans = self._generate_plans(activities)
        doc.add_paragraph(plans)
        
        # 保存
        filename = f'社区活动汇总报告_{datetime.now().strftime("%Y%m%d")}.docx'
        report_path = os.path.join(self.output_dir, filename)
        doc.save(report_path)
        
        return report_path
    
    def _get_report_period(self, activities: List[Dict[str, Any]]) -> str:
        """获取报告周期"""
        if not activities:
            return datetime.now().strftime('%Y年%m月')
        
        dates = []
        for activity in activities:
            date_str = activity.get('activity_date', '')
            if date_str:
                try:
                    # 尝试解析 "2026.6.26" 格式
                    if '.' in date_str:
                        parts = date_str.split('.')
                        if len(parts) >= 2:
                            year = int(parts[0]) if len(parts[0]) == 4 else datetime.now().year
                            month = int(parts[1])
                            dates.append((year, month))
                except:
                    pass
        
        if dates:
            dates.sort()
            start = dates[0]
            end = dates[-1]
            if start == end:
                return f'{start[0]}年{start[1]}月'
            else:
                return f'{start[0]}年{start[1]}月 - {end[0]}年{end[1]}月'
        
        return datetime.now().strftime('%Y年%m月')
    
    def _format_date(self, date_str: str) -> str:
        """格式化日期"""
        if not date_str:
            return ''
        
        try:
            if '.' in date_str:
                parts = date_str.split('.')
                if len(parts) == 3:
                    return f'{parts[0]}年{parts[1]}月{parts[2]}日'
                elif len(parts) == 2:
                    return f'{datetime.now().year}年{parts[0]}月{parts[1]}日'
        except:
            pass
        
        return date_str
    
    def _generate_overview(self, activities: List[Dict[str, Any]]) -> str:
        """生成总体概况"""
        if not activities:
            return '暂无活动记录。'
        
        total = len(activities)
        
        # 统计总参与人数
        total_participants = 0
        for activity in activities:
            participants = activity.get('actual_participants', activity.get('estimated_participants', ''))
            if participants:
                # 尝试提取数字
                import re
                nums = re.findall(r'\d+', str(participants))
                if nums:
                    total_participants += sum(int(n) for n in nums)
        
        overview = f"""本周期内共举办社区活动 {total} 场，累计参与人数约 {total_participants} 人次。

活动类型涵盖感统训练、亲子互动、角色扮演等多种形式，服务对象包括婴幼儿及家庭、社区居民等群体。各项活动均按照预定方案有序开展，现场氛围良好，参与者反馈积极。
"""
        
        return overview
    
    def _generate_statistics(self, activities: List[Dict[str, Any]]) -> str:
        """生成统计数据"""
        if not activities:
            return ''
        
        # 按主题分类统计
        themes = {}
        for activity in activities:
            theme = activity.get('activity_theme', '未分类')
            if theme in themes:
                themes[theme] += 1
            else:
                themes[theme] = 1
        
        stats_lines = [f'活动总数：{len(activities)}场', '']
        
        if themes:
            stats_lines.append('活动分类：')
            for theme, count in themes.items():
                stats_lines.append(f'  • {theme}：{count}场')
        
        stats_lines.append('')
        
        # 参与人数统计
        participant_counts = []
        for activity in activities:
            participants = activity.get('actual_participants', activity.get('estimated_participants', ''))
            if participants:
                import re
                nums = re.findall(r'\d+', str(participants))
                if nums:
                    participant_counts.extend([int(n) for n in nums])
        
        if participant_counts:
            stats_lines.append(f'参与人数统计：')
            stats_lines.append(f'  • 单场最多：{max(participant_counts)}人')
            stats_lines.append(f'  • 单场最少：{min(participant_counts)}人')
            stats_lines.append(f'  • 平均参与：{sum(participant_counts)//len(participant_counts)}人')
        
        return '\n'.join(stats_lines)
    
    def _generate_analysis(self, activities: List[Dict[str, Any]]) -> str:
        """生成复盘分析"""
        if not activities:
            return '暂无数据可供分析。'
        
        # 收集亮点和问题
        all_highlights = []
        all_problems = []
        all_improvements = []
        
        for activity in activities:
            highlights = activity.get('highlights', '')
            if highlights:
                all_highlights.append(highlights)
            
            problems = activity.get('problems', '')
            if problems:
                all_problems.append(problems)
            
            improvements = activity.get('improvements', '')
            if improvements:
                all_improvements.append(improvements)
        
        analysis_parts = []
        
        # 共性亮点
        if all_highlights:
            analysis_parts.append('（一）共性亮点')
            for i, highlight in enumerate(all_highlights[:3], 1):
                analysis_parts.append(f'{i}. {highlight}')
            analysis_parts.append('')
        
        # 存在问题
        if all_problems:
            analysis_parts.append('（二）存在问题')
            for i, problem in enumerate(all_problems[:3], 1):
                analysis_parts.append(f'{i}. {problem}')
            analysis_parts.append('')
        
        # 改进方向
        if all_improvements:
            analysis_parts.append('（三）改进方向')
            for i, improvement in enumerate(all_improvements[:3], 1):
                analysis_parts.append(f'{i}. {improvement}')
            analysis_parts.append('')
        
        if not analysis_parts:
            analysis_parts.append('各项活动均按计划顺利开展，暂未发现明显问题。建议持续关注参与者反馈，不断优化活动内容和形式。')
        
        return '\n'.join(analysis_parts)
    
    def _generate_plans(self, activities: List[Dict[str, Any]]) -> str:
        """生成下一步计划"""
        if not activities:
            return '暂无计划。'
        
        # 使用AI生成计划
        context = self._build_context(activities)
        
        system_prompt = "你是一位社区活动策划专家，擅长根据已有活动经验制定下一步计划。"
        
        prompt = f"""
请根据以下社区活动记录，撰写下一步工作计划（200字左右）：

{context}

要求：
1. 基于已有活动的经验和问题
2. 提出1-3个具体改进方向
3. 建议下次活动的主题或形式
4. 语言简洁，适合工作汇报
"""
        
        result = self.llm.chat(prompt, system_prompt)
        
        if not result or len(result) < 50:
            # 降级方案
            plans = """1. 持续优化活动流程，提升参与者体验
2. 根据居民反馈，丰富活动内容形式
3. 加强活动宣传，扩大参与覆盖面
4. 完善活动记录，建立活动档案库"""
            return plans
        
        return result
    
    def _build_context(self, activities: List[Dict[str, Any]]) -> str:
        """构建上下文信息"""
        lines = []
        
        for i, activity in enumerate(activities, 1):
            lines.append(f"活动{i}：{activity.get('activity_theme', '')}")
            lines.append(f"  时间：{activity.get('activity_date', '')}")
            lines.append(f"  参与：{activity.get('actual_participants', activity.get('estimated_participants', ''))}")
            
            if activity.get('highlights'):
                lines.append(f"  亮点：{activity['highlights']}")
            if activity.get('problems'):
                lines.append(f"  问题：{activity['problems']}")
            if activity.get('improvements'):
                lines.append(f"  改进：{activity['improvements']}")
            
            lines.append('')
        
        return '\n'.join(lines)
